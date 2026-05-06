# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path(r"C:\Users\86157\Desktop\UniAppPencil")
DOC_DIR = ROOT / "毕设文档"
BACKUP = DOC_DIR / "毕业论文模板_原始备份.docx"
TEMPLATE = DOC_DIR / "毕业论文模板.docx"
OUTPUT = DOC_DIR / "毕业论文_徐超_格式修正版.docx"
DESKTOP = Path(r"C:\Users\86157\Desktop\毕业论文_徐超_格式修正版.docx")

TITLE_CN = "基于 Node.js 与 JavaScript 实现购物小程序的设计与实现"
TITLE_EN = "Design and Implementation of a Shopping Mini-Program Based on Node.js and JavaScript"


TOC_LINES = [
    ("1. 绪论", "1"),
    ("1.1  选题背景", "1"),
    ("1.2  研究现状", "1"),
    ("1.3  研究意义", "2"),
    ("1.4  论文结构", "3"),
    ("2. 系统开发技术与环境", "4"),
    ("2.1  开发工具", "4"),
    ("2.2  开发技术", "4"),
    ("2.3  系统运行环境", "5"),
    ("3. 系统分析", "6"),
    ("3.1  可行性分析", "6"),
    ("3.2  需求分析", "7"),
    ("3.2.1  功能需求分析", "7"),
    ("3.2.2  性能需求分析", "10"),
    ("3.2.3  安全需求分析", "11"),
    ("3.2.4  兼容性需求分析", "11"),
    ("4. 系统设计", "12"),
    ("4.1  系统总体功能设计", "12"),
    ("4.2  系统架构设计", "13"),
    ("4.3  系统业务流程设计", "14"),
    ("4.4  系统数据库设计", "15"),
    ("4.4.1  数据库需求分析", "15"),
    ("4.4.2  数据库概念结构设计", "16"),
    ("4.4.3  数据表设计", "16"),
    ("5. 系统功能实现", "18"),
    ("5.1  用户端功能实现", "18"),
    ("5.1.1  用户登录与身份识别", "18"),
    ("5.1.2  首页商品展示与分类浏览", "19"),
    ("5.1.3  商品搜索与筛选", "20"),
    ("5.1.4  商品详情浏览", "21"),
    ("5.1.5  购物车管理", "23"),
    ("5.1.6  订单确认与订单提交", "24"),
    ("5.1.7  地址管理", "25"),
    ("5.1.8  售后申请", "26"),
    ("5.1.9  用户反馈提交", "27"),
    ("5.1.10  AI 智能客服聊天", "28"),
    ("5.1.11  个人中心与资料管理", "29"),
    ("5.2  管理端功能实现", "30"),
    ("5.2.1  后台登录与权限控制", "30"),
    ("5.2.2  商品管理", "31"),
    ("5.2.3  订单管理", "32"),
    ("5.2.4  售后管理", "33"),
    ("5.2.5  反馈管理", "33"),
    ("5.2.6  系统数据统计与概览", "34"),
    ("6. 系统测试", "35"),
    ("6.1  测试目标与范围", "35"),
    ("6.2  测试环境与工具", "35"),
    ("6.3  测试方法与策略", "36"),
    ("6.4  功能测试", "36"),
    ("6.4.1  核心功能测试用例表", "36"),
    ("6.4.2  测试结论分析", "37"),
    ("6.5  非功能测试", "38"),
    ("6.6  测试结果分析与总结", "39"),
    ("结论", "40"),
    ("致谢", "41"),
    ("参考文献", "42"),
]


def add_para(doc, text="", style=None, align=None):
    p = doc.add_paragraph(text, style=style)
    if align is not None:
        p.alignment = align
    return p


def add_chapter(doc, title):
    add_para(doc, title, style="Heading 1", align=WD_ALIGN_PARAGRAPH.CENTER)


def add_section(doc, title):
    add_para(doc, title, style="Heading 2")


def add_subsection(doc, title):
    add_para(doc, title, style="Heading 3")


def add_figure_placeholder(doc, caption):
    p = add_para(doc, "【此处预留图片位置】", align=WD_ALIGN_PARAGRAPH.CENTER)
    p.runs[0].font.size = Pt(11)
    add_para(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER)


def add_table(doc, headers, rows, caption):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    add_para(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER)


def update_cover(doc):
    if len(doc.tables) >= 2:
        t0 = doc.tables[0]
        t0.cell(0, 1).text = TITLE_CN
        t0.cell(1, 1).text = TITLE_EN
        t0.cell(2, 1).text = "徐超"
        t0.cell(2, 3).text = "工学"
        t0.cell(3, 1).text = "2022213244"
        t0.cell(4, 1).text = "软件工程"
        t0.cell(5, 1).text = "软件学院"
        t0.cell(6, 1).text = "郑勇明"
        t0.cell(6, 3).text = "讲师"

        t1 = doc.tables[1]
        t1.cell(0, 1).text = "Xu Chao"
        t1.cell(1, 1).text = "Engineering"
        t1.cell(2, 1).text = "2022213244"
        t1.cell(3, 1).text = "Software Engineering"
        t1.cell(4, 1).text = "School of Software"
        t1.cell(5, 1).text = "Zheng Yongming"
        t1.cell(6, 1).text = "Lecturer"

    fixed = {
        11: "二〇二六 年 五 月 六 日",
        18: TITLE_EN,
        33: "Date: 2026-05-06",
        43: "                              签字日期： 2026年 5月 6日",
        48: "                              签字日期： 2026年 5月 6日",
    }
    for idx, text in fixed.items():
        if idx < len(doc.paragraphs):
            doc.paragraphs[idx].text = text


def remove_old_body(doc):
    start = doc.paragraphs[52]._element
    parent = start.getparent()
    children = list(parent)
    start_idx = children.index(start)
    for child in children[start_idx:]:
        parent.remove(child)


def add_front_matter(doc):
    add_para(doc, "摘  要", style="Heading 1", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "随着移动互联网和微信生态的持续发展，小程序因其即用即走、开发成本较低、使用门槛较小等特点，逐渐成为中小型电商系统的重要载体。针对传统购物小程序在商品展示、订单管理、用户交互和智能咨询方面存在的功能碎片化、体验割裂以及客服响应效率不足等问题，本文设计并实现了一套基于 Node.js 与 JavaScript 的智能购物小程序系统。系统采用前后端分离的开发模式，前端基于 uni-app 开发微信小程序界面，后端基于 Express 构建 RESTful API 服务，数据库采用 MySQL 进行数据持久化存储，并结合第三方大模型能力构建 AI 智能客服模块。")
    add_para(doc, "系统围绕“浏览商品、智能咨询、加入购物车、下单支付、订单追踪、售后处理”的完整业务闭环展开设计，重点实现了用户登录与信息管理、商品分类展示与搜索、商品详情浏览、购物车管理、订单流转、AI 聊天咨询以及后台管理等功能。与传统电商小程序相比，本系统在商品详情页中引入上下文感知的智能问答能力，使用户可以围绕商品属性、使用方式、物流与售后等问题进行自然语言咨询，从而提升了交互效率与用户体验。")
    add_para(doc, "在系统实现过程中，本文重点讨论了前后端接口设计、数据库表结构设计、事务处理机制以及 AI 接口调用的稳定性控制方案。测试结果表明，系统能够较好地满足基本购物场景的功能需求，具备一定的可扩展性和实际应用价值。")
    add_para(doc, "关键词：微信小程序；Node.js；Express；MySQL；智能客服；电商系统")
    doc.add_page_break()

    add_para(doc, "ABSTRACT", style="Heading 1", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "With the continuous development of mobile Internet and the WeChat ecosystem, mini programs have become an important carrier for small and medium-sized e-commerce systems due to their lightweight usage and relatively low development cost. To address the problems of fragmented functions, inconsistent user experience, and inefficient customer service in traditional shopping mini programs, this paper designs and implements an intelligent shopping mini program based on Node.js and JavaScript.")
    add_para(doc, "The system focuses on the complete business workflow of product browsing, intelligent consultation, cart management, order placement, order tracking, and after-sales handling. It implements user authentication and profile management, product classification and search, product details browsing, shopping cart management, order lifecycle control, AI chat consultation, and backend administration.")
    add_para(doc, "This paper also discusses front-end and back-end API design, database schema design, transaction processing, and stability control for AI API calls. Test results show that the system can satisfy the basic functional requirements of shopping scenarios and has good scalability and practical value.")
    add_para(doc, "Key words: WeChat Mini Program; Node.js; Express; MySQL; Intelligent Customer Service; E-commerce System")
    doc.add_page_break()

    add_para(doc, "目  录", style="Heading 1", align=WD_ALIGN_PARAGRAPH.CENTER)
    for title, page in TOC_LINES:
        p = add_para(doc)
        if title.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "结论", "致谢", "参考文献")) and title.count(".") <= 1:
            p.paragraph_format.left_indent = Pt(0)
        else:
            p.paragraph_format.left_indent = Pt(14)
        p.add_run(title)
        p.add_run("\t" + page)
    doc.add_page_break()


def add_body(doc):
    add_chapter(doc, "第1章  绪论")
    add_section(doc, "1.1  选题背景")
    add_para(doc, "随着移动互联网与即时通信技术的深度融合，微信小程序凭借“即用即走、无需安装”的特点，已经成为轻量级电商服务的重要载体。相比传统 App，小程序具有使用门槛低、传播路径短、维护成本低等优势，尤其适合校园商城、社区团购、个体商户和中小型店铺等场景。")
    add_para(doc, "当前不少购物小程序仍然存在功能单一、交互割裂、客服响应滞后等问题。用户在浏览商品后，往往还需要通过人工客服确认商品细节、物流规则和售后政策，购买链路不够连贯。与此同时，大模型技术的发展为自然语言交互提供了新的技术路径，将 AI 能力接入购物小程序，有助于提升用户咨询效率和系统智能化水平。")
    add_section(doc, "1.2  研究现状")
    add_para(doc, "国内外关于电商系统、小程序开发和智能客服的研究较多。传统电商系统主要关注商品展示、订单管理、支付处理和用户评价等业务流程；微信小程序进一步降低了用户使用门槛，使中小型商户能够以较低成本接入移动端经营场景。")
    add_para(doc, "在智能客服方面，早期系统多依赖规则库和关键词匹配，能够处理固定问题，但对于商品上下文和复杂语义的理解能力不足。近年来，DeepSeek、OpenAI 等大模型接口逐渐成熟，为电商场景中的智能问答、导购推荐和售后咨询提供了更灵活的实现方式。")
    add_section(doc, "1.3  研究意义")
    add_para(doc, "本文设计并实现的购物小程序系统，将 uni-app 小程序前端、Node.js 后端、MySQL 数据库和 AI 智能客服能力结合起来，形成“浏览—咨询—加购—下单—售后”的完整业务闭环。该系统不仅能够验证轻量级前后端分离架构在购物场景中的可行性，也能探索 AI 能力在移动端电商服务中的应用价值。")
    add_section(doc, "1.4  论文结构")
    add_para(doc, "本文共分为六章。第1章介绍选题背景、研究现状和研究意义；第2章介绍系统开发技术与环境；第3章进行可行性分析和需求分析；第4章阐述系统总体设计、架构设计和数据库设计；第5章说明用户端、管理端及核心业务功能实现；第6章对系统进行测试和结果分析。")

    add_chapter(doc, "第2章  系统开发技术与环境")
    add_section(doc, "2.1  开发工具")
    add_para(doc, "本系统在开发过程中主要使用 Visual Studio Code 进行前端页面和后端代码编写，使用微信小程序开发工具或 H5 调试环境进行页面预览，使用 MySQL 数据库管理工具进行数据表维护和数据检查。项目代码通过前后端分离方式组织，便于模块开发和接口联调。")
    add_section(doc, "2.2  开发技术")
    add_para(doc, "前端采用 uni-app 框架开发微信小程序界面，并使用 Vue 语法组织页面逻辑。系统通过 Pinia 管理用户状态和购物车状态，通过 services 层统一封装接口请求，页面层不直接拼接接口地址。")
    add_para(doc, "后端采用 Node.js 与 Express 构建 RESTful API 服务。项目按照 auth、product、cart、order、chat、address、afterSale、feedback、admin 等业务域划分模块，并采用 routes、controller、service、repository 的分层结构。数据库采用 MySQL 保存用户、商品、购物车、订单、地址、售后、反馈和聊天记录等数据。")
    add_para(doc, "AI 智能客服模块由后端统一封装大模型接口，前端只调用本系统聊天接口，不直接接触 API 密钥。该设计既保护了密钥安全，也方便后端统一处理超时、异常和降级提示。")
    add_section(doc, "2.3  系统运行环境")
    add_para(doc, "系统运行环境包括 Windows 本地开发环境、Node.js 运行环境、MySQL 数据库环境和小程序调试环境。后端服务通过 Express 启动，前端通过 uni-app 运行并请求后端接口，数据库负责持久化存储业务数据。")

    add_chapter(doc, "第3章  系统分析")
    add_section(doc, "3.1  可行性分析")
    add_para(doc, "在技术可行性方面，uni-app、Node.js、Express 和 MySQL 均为成熟技术，能够支撑本系统的前端页面、后端接口和数据存储需求。在经济可行性方面，系统采用的主要技术均为开源方案，本地开发和测试成本较低。在运行可行性方面，系统围绕普通用户和管理员两类角色设计，业务路径清晰，用户学习成本较低。")
    add_section(doc, "3.2  需求分析")
    add_subsection(doc, "3.2.1  功能需求分析")
    add_para(doc, "普通用户需要完成登录、商品浏览、商品搜索、商品详情查看、购物车管理、订单确认、地址管理、售后申请、反馈提交、AI 智能客服咨询和个人中心管理等操作。管理员需要完成后台登录、商品管理、订单管理、售后管理、反馈管理和系统数据统计等操作。")
    add_figure_placeholder(doc, "图3-1 用户端功能结构图")
    add_figure_placeholder(doc, "图3-2 管理端功能结构图")
    add_subsection(doc, "3.2.2  性能需求分析")
    add_para(doc, "系统应保证首页商品列表、搜索结果、商品详情、购物车和订单等高频接口具有较好的响应速度。订单创建场景需要保证事务一致性，避免库存扣减、订单生成和购物车清理之间出现数据不一致。AI 接口可能存在网络延迟或第三方服务异常，因此需要设置超时和降级提示。")
    add_subsection(doc, "3.2.3  安全需求分析")
    add_para(doc, "系统涉及用户资料、地址、订单和聊天记录等私有数据，因此必须进行身份认证和权限控制。后端通过 requireAuth 保护用户私有接口，通过 requireAdmin 保护后台接口。AI 接口密钥仅保存在后端环境变量中，前端无法直接获取。")
    add_subsection(doc, "3.2.4  兼容性需求分析")
    add_para(doc, "系统前端应适配主流微信客户端和常见移动设备。页面使用统一组件和响应式布局方式，保证首页、搜索、购物车、订单和个人中心等主要页面在不同设备上的显示效果一致。")

    add_chapter(doc, "第4章  系统设计")
    add_section(doc, "4.1  系统总体功能设计")
    add_para(doc, "系统采用普通用户端和管理员端双角色设计。用户端围绕商品浏览、搜索、咨询、加购、下单、售后和反馈展开；管理端围绕商品维护、订单处理、售后审批、反馈查看和运营概览展开。")
    add_figure_placeholder(doc, "图4-1 系统总体功能模块图")
    add_section(doc, "4.2  系统架构设计")
    add_para(doc, "系统采用前后端分离架构，整体分为前端展示层、后端服务层、数据存储层和 AI 能力层。前端基于 uni-app 构建小程序页面，后端基于 Express 提供 RESTful API，MySQL 负责持久化存储，AI 能力通过后端代理调用第三方大模型接口实现。")
    add_figure_placeholder(doc, "图4-2 系统总体架构图")
    add_section(doc, "4.3  系统业务流程设计")
    add_para(doc, "用户下单流程为：用户浏览商品并进入详情页，选择规格和数量后加入购物车，进入确认订单页选择地址并提交订单，后端在事务中生成订单、写入订单明细并清理购物车。AI 咨询流程为：用户从商品详情页进入聊天页，前端携带商品 ID 创建会话，后端拼接商品上下文调用 AI 接口并保存问答记录。")
    add_figure_placeholder(doc, "图4-3 下单业务流程图")
    add_figure_placeholder(doc, "图4-4 AI 咨询业务流程图")
    add_section(doc, "4.4  系统数据库设计")
    add_subsection(doc, "4.4.1  数据库需求分析")
    add_para(doc, "系统数据库需要保存用户、商品、商品图集、购物车、地址、订单、订单明细、售后、反馈、聊天会话和聊天消息等数据。这些数据共同支撑完整购物业务闭环。")
    add_subsection(doc, "4.4.2  数据库概念结构设计")
    add_para(doc, "用户与地址、购物车、订单、售后、反馈和聊天会话之间是一对多关系；商品与图集、购物车条目、订单明细和反馈之间存在关联关系；订单与订单明细、售后申请之间存在业务关联。")
    add_figure_placeholder(doc, "图4-5 数据库 E-R 图")
    add_subsection(doc, "4.4.3  数据表设计")
    add_table(doc, ["表名", "主要作用", "关联模块"], [
        ["users", "保存用户身份、角色和资料", "认证、个人中心、后台权限"],
        ["products", "保存商品基础信息", "首页、搜索、详情、后台商品管理"],
        ["product_galleries", "保存商品图集", "商品详情、后台商品表单"],
        ["cart_items", "保存购物车条目", "购物车、订单确认"],
        ["addresses", "保存收货地址", "地址管理、订单确认"],
        ["orders", "保存订单主信息", "订单列表、后台订单管理"],
        ["order_items", "保存订单商品明细", "订单详情、售后"],
        ["after_sales", "保存售后申请", "售后申请、后台审批"],
        ["feedbacks", "保存用户反馈", "反馈提交、后台反馈管理"],
        ["chat_sessions", "保存 AI 会话", "AI 聊天、历史记录"],
        ["chat_messages", "保存聊天消息", "AI 聊天、历史记录"],
    ], "表4-1 系统主要数据表")

    add_chapter(doc, "第5章  系统功能实现")
    add_section(doc, "5.1  用户端功能实现")
    user_sections = [
        ("5.1.1  用户登录与身份识别", "登录页面通过 authService 向后端提交登录请求，后端返回用户资料和角色信息。前端使用 Pinia 的 user store 维护登录状态，个人中心根据身份展示普通用户入口或后台管理入口。"),
        ("5.1.2  首页商品展示与分类浏览", "首页展示 Banner、分类入口、热销商品列表和 AI 导购入口。前端调用 productService 获取商品数据，并通过 ProductCard 组件展示商品信息。"),
        ("5.1.3  商品搜索与筛选", "搜索页支持关键词、分类和排序条件。前端将查询条件转为 query 参数，后端根据条件动态查询商品并返回列表。"),
        ("5.1.4  商品详情浏览", "商品详情页展示商品主图、图集、价格、库存、规格、服务标签和底部购买操作。用户可在详情页加入购物车、立即购买或发起 AI 咨询。"),
        ("5.1.5  购物车管理", "购物车页支持商品数量修改、删除和金额汇总。前端通过 cartService 与后端同步数据，后端校验商品存在性和库存。"),
        ("5.1.6  订单确认与订单提交", "确认订单页展示收货地址、商品列表、金额和备注。后端创建订单时通过事务保证订单写入、订单明细写入和购物车清理的一致性。"),
        ("5.1.7  地址管理", "地址页支持新增、编辑、删除和默认地址设置。地址数据与当前用户绑定，订单确认时可自动读取默认地址。"),
        ("5.1.8  售后申请", "用户可从订单列表进入售后申请页，填写售后原因和说明。后端保存售后记录，管理员可在后台更新处理状态。"),
        ("5.1.9  用户反馈提交", "反馈模块用于收集用户对商品或系统体验的意见。后端保存反馈内容，后台可查看并更新处理状态。"),
        ("5.1.10  AI 智能客服聊天", "AI 聊天页支持商品上下文咨询。用户从商品详情页进入时，系统携带商品 ID 创建会话，后端调用 AI 接口生成回答并保存消息记录。"),
        ("5.1.11  个人中心与资料管理", "个人中心展示用户昵称、头像、订单、地址、售后、反馈和 AI 历史入口。用户可编辑头像和昵称等资料。"),
    ]
    for title, text in user_sections:
        add_subsection(doc, title)
        add_para(doc, text)
        if title in ("5.1.2  首页商品展示与分类浏览", "5.1.4  商品详情浏览", "5.1.10  AI 智能客服聊天"):
            add_figure_placeholder(doc, "图" + title.split()[0] + " 功能运行界面图")

    add_section(doc, "5.2  管理端功能实现")
    admin_sections = [
        ("5.2.1  后台登录与权限控制", "后台管理接口由 requireAdmin 中间件保护，只有管理员角色可以访问。前端也会根据用户角色控制后台入口显示。"),
        ("5.2.2  商品管理", "管理员可以新增、编辑、上下架商品，并维护商品图集、库存、价格和标签等信息。"),
        ("5.2.3  订单管理", "管理员可以查看订单列表和订单详情，并更新订单状态和物流信息，使用户端同步展示订单进度。"),
        ("5.2.4  售后管理", "管理员可查看用户提交的售后申请，依据原因和订单信息更新处理状态。"),
        ("5.2.5  反馈管理", "管理员可查看用户反馈内容，并将反馈状态标记为待处理或已处理。"),
        ("5.2.6  系统数据统计与概览", "后台首页汇总订单、商品、售后、反馈和 AI 调用数据，为管理员提供系统运行概览。"),
    ]
    for title, text in admin_sections:
        add_subsection(doc, title)
        add_para(doc, text)

    add_chapter(doc, "第6章  系统测试")
    add_section(doc, "6.1  测试目标与范围")
    add_para(doc, "系统测试目标是验证用户端、管理端和后端接口是否满足设计要求。测试范围覆盖登录、商品浏览、搜索、购物车、订单、地址、售后、反馈、AI 聊天和后台管理等核心功能。")
    add_section(doc, "6.2  测试环境与工具")
    add_para(doc, "测试环境为 Windows 本地开发环境，前端使用 uni-app 调试，后端使用 Node.js 启动 Express 服务，数据库使用 MySQL。")
    add_section(doc, "6.3  测试方法与策略")
    add_para(doc, "测试采用功能测试、边界测试、权限测试和异常测试相结合的方式。对于订单和购物车等核心链路，重点验证数据一致性；对于 AI 模块，重点验证接口异常时的降级提示。")
    add_section(doc, "6.4  功能测试")
    add_subsection(doc, "6.4.1  核心功能测试用例表")
    add_table(doc, ["用例编号", "测试项", "前置条件", "操作步骤", "预期结果", "实测结果"], [
        ["TC-01", "用户登录", "存在有效用户", "提交登录信息", "登录成功并保存状态", "通过"],
        ["TC-02", "商品浏览", "商品数据存在", "打开首页", "商品列表正常展示", "通过"],
        ["TC-03", "购物车管理", "用户已登录", "加入商品并修改数量", "数量和金额同步变化", "通过"],
        ["TC-04", "订单提交", "购物车有商品", "提交订单", "订单创建成功", "通过"],
        ["TC-05", "AI 咨询", "用户已登录", "发送商品相关问题", "返回智能回答", "通过"],
        ["TC-06", "后台管理", "管理员已登录", "进入商品管理页", "可查看和编辑商品", "通过"],
    ], "表6-1 核心功能测试用例表")
    add_subsection(doc, "6.4.2  测试结论分析")
    add_para(doc, "测试结果表明，系统核心功能能够正常运行，前后端接口返回稳定，用户端和管理端的主要业务流程符合预期。")
    add_section(doc, "6.5  非功能测试")
    add_para(doc, "非功能测试主要包括兼容性测试、响应时效测试、安全机制测试和运行稳定性观察。系统在常见调试环境下能够保持页面展示和操作流程稳定，权限接口能够正确拦截未授权访问。")
    add_section(doc, "6.6  测试结果分析与总结")
    add_para(doc, "综合测试结果可知，本系统能够满足毕业设计对购物小程序的基本功能和稳定性要求。后续可继续扩展真实支付、评价体系、推荐算法和日志监控等能力。")

    add_chapter(doc, "结论")
    add_para(doc, "本文完成了一套基于 Node.js 与 JavaScript 的购物小程序设计与实现。系统以 uni-app 构建前端页面，以 Express 构建后端服务，以 MySQL 保存核心业务数据，并结合 AI 接口实现智能客服功能。系统覆盖了登录、商品展示、搜索、详情、购物车、订单、地址、售后、反馈、AI 聊天和后台管理等模块，形成较完整的购物业务闭环。")
    add_para(doc, "从实现结果看，该系统具有模块划分清晰、接口结构稳定、业务流程完整和可扩展性较好的特点。后续可继续完善支付、评价、推荐和运营统计等功能，使系统更加贴近真实电商应用场景。")

    add_chapter(doc, "致  谢")
    add_para(doc, "在本论文和系统实现过程中，感谢指导教师郑勇明老师在选题方向、系统设计和论文写作方面给予的指导，也感谢在项目开发过程中提供帮助的同学和朋友。通过本次毕业设计，本人进一步加深了对微信小程序开发、Node.js 后端构建、MySQL 数据库设计和 AI 接口集成的理解，提升了综合工程实践能力。")

    add_chapter(doc, "参 考 文 献")
    refs = [
        "陈宇收，饶宏博，王英明，等. 基于 JWT 的前后端分离程序设计研究[J]. 电脑编程技巧与维护，2019(09):11-12.",
        "钟小平. Node.js 开发实战教程[M]. 北京：人民邮电出版社，2020.",
        "黎青霞. Node.js 在 Web 开发中的应用研究[J]. 信息记录材料，2024,25(10):91-93+96.",
        "杜雨荃，王晓菊，田立勤. 基于微信小程序的网上购物系统的设计与实现[J].",
        "马静. 基于微信小程序的购物商城系统的设计与实现[J]. 微型电脑应用，2021(03).",
        "李哲，周灵. 微信小程序的架构与开发浅析[J]. 福建电脑，2019(02).",
    ]
    for i, ref in enumerate(refs, 1):
        add_para(doc, f"[{i}] {ref}")


def main():
    source = BACKUP if BACKUP.exists() else TEMPLATE
    doc = Document(str(source))
    update_cover(doc)
    remove_old_body(doc)
    add_front_matter(doc)
    add_body(doc)
    doc.save(str(OUTPUT))
    doc.save(str(TEMPLATE))
    try:
        doc.save(str(DESKTOP))
    except PermissionError:
        pass
    print(OUTPUT)


if __name__ == "__main__":
    main()
