# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document


PATH = Path(r"C:\Users\86157\Desktop\UniAppPencil\毕设文档\毕业论文_徐超_格式修正版.docx")


def main() -> None:
    doc = Document(str(PATH))
    text = "\n".join(par.text for par in doc.paragraphs)
    print("paragraphs", len(doc.paragraphs))
    print("tables", len(doc.tables))
    print("chars", len(text))
    for needle in ["SpringBoot", "宠物", "健身", "TODO", "TBD", "农场", "Java Spring"]:
        print(f"{needle}={needle in text}")
    print("head")
    for idx, par in enumerate(doc.paragraphs[:45]):
        if par.text.strip():
            print(idx, par.text[:120])


if __name__ == "__main__":
    main()
