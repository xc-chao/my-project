# -*- coding: utf-8 -*-
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.shared import Pt, Cm


ROOT = Path(r"C:\Users\86157\Desktop\UniAppPencil")
DOC_DIR = ROOT / "毕设文档"
BACKUP = DOC_DIR / "毕业论文模板_原始备份.docx"
FALLBACK_TEMPLATE = DOC_DIR / "毕业论文模板.docx"
OUTPUT = DOC_DIR / "毕业论文_徐超_格式修正版.docx"
DESKTOP = Path(r"C:\Users\86157\Desktop\毕业论文_徐超_格式修正版.docx")
DESKTOP_ALT = Path(r"C:\Users\86157\Desktop\毕业论文_徐超_格式修正版_新版.docx")

TITLE_CN = "基于 Node.js 与 JavaScript 实现购物小程序的设计与实现"
TITLE_EN = "Design and Implementation of a Shopping Mini-Program Based on Node.js and JavaScript"


TOC_LINES = [
    ("1. 绪论", "1"),
    ("1.1  选题的背景", "1"),
    ("1.2  研究现状", "1"),
    ("1.3  研究意义", "2"),
    ("1.4  论文结构", "3"),
    ("2. 系统开发技术与环境", "4"),
    ("2.1  开发工具", "4"),
    ("2.2  开发技术", "4"),
    ("2.2.1  uni-app 小程序开发技术", "4"),
    ("2.2.2  Vue 与 Pinia 状态管理", "5"),
    ("2.2.3  Node.js 与 Express 后端技术", "5"),
    ("2.2.4  MySQL 数据库技术", "6"),
    ("2.2.5  AI 智能客服接口技术", "6"),
    ("2.3  系统运行环境", "6"),
    ("3. 系统分析", "7"),
    ("3.1  可行性分析", "7"),
    ("3.1.1  技术可行性", "7"),
    ("3.1.2  经济可行性", "8"),
    ("3.1.3  运行可行性", "8"),
    ("3.2  需求分析", "9"),
    ("3.2.1  功能需求分析", "9"),
    ("3.2.2  性能需求分析", "12"),
    ("3.2.3  安全需求分析", "13"),
    ("3.2.4  兼容性需求分析", "13"),
    ("4. 系统设计", "14"),
    ("4.1  系统总体功能设计", "14"),
    ("4.2  系统架构设计", "15"),
    ("4.3  系统业务流程设计", "16"),
    ("4.4  系统数据库设计", "17"),
    ("4.4.1  数据库需求分析", "17"),
    ("4.4.2  数据库概念结构设计", "18"),
    ("4.4.3  数据表设计", "18"),
    ("5. 系统功能实现", "20"),
    ("5.1  用户端功能的实现", "20"),
    ("5.1.1  用户登录注册", "20"),
    ("5.1.2  首页商品展示与搜索", "21"),
    ("5.1.3  商品详情与反馈", "23"),
    ("5.1.4  购物车管理", "25"),
    ("5.1.5  订单确认与地址管理", "26"),
    ("5.1.6  售后申请", "28"),
    ("5.1.7  AI 智能客服聊天", "29"),
    ("5.1.8  个人中心与资料管理", "31"),
    ("5.2  管理端功能实现", "32"),
    ("5.2.1  后台登录与权限控制", "32"),
    ("5.2.2  商品管理", "33"),
    ("5.2.3  订单管理", "34"),
    ("5.2.4  售后与反馈管理", "35"),
    ("5.2.5  系统数据统计与概览", "36"),
    ("5.3  后端核心模块实现", "37"),
    ("6. 系统测试", "39"),
    ("6.1  测试目标与范围", "39"),
    ("6.2  测试环境与工具", "39"),
    ("6.3  测试方法与策略", "40"),
    ("6.4  功能测试", "41"),
    ("6.4.1  核心功能测试用例表", "41"),
    ("6.4.2  测试结论分析", "42"),
    ("6.5  非功能测试", "43"),
    ("6.6  测试结果分析与总结", "44"),
    ("结论", "45"),
    ("致  谢", "46"),
    ("参考文献", "47"),
]


def set_doc_defaults(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    for name in ("Heading 1", "Heading 2", "Heading 3"):
        if name in styles:
            style = styles[name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_para(doc: Document, text: str = "", style: str | None = None, align=None):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(12)
    p.paragraph_format.first_line_indent = Pt(24) if style is None and text else None
    p.paragraph_format.line_spacing = 1.5
    if align is not None:
        p.alignment = align
    return p


def add_heading(doc: Document, text: str, level: int) -> None:
    if level == 1:
        p = doc.add_paragraph(style="Heading 1")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        p = doc.add_paragraph(style="Heading 2")
    else:
        p = doc.add_paragraph(style="Heading 3")
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(15 if level == 1 else 14 if level == 2 else 12)
    r.bold = True
    p.paragraph_format.line_spacing = 1.5


def add_figure_placeholder(doc: Document, caption: str) -> None:
    p = add_para(doc, "【此处预留图片位置】", align=WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.first_line_indent = None
    add_para(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER).paragraph_format.first_line_indent = None


def add_table(doc: Document, headers, rows, caption: str) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    add_para(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER).paragraph_format.first_line_indent = None


def update_cover(doc: Document) -> None:
    if len(doc.tables) >= 2:
        t0 = doc.tables[0]
        replacements = {
            (0, 1): TITLE_CN,
            (1, 1): TITLE_EN,
            (2, 1): "徐超",
            (2, 3): "工学",
            (3, 1): "2022213244",
            (4, 1): "软件工程",
            (5, 1): "软件学院",
            (6, 1): "郑勇明",
            (6, 3): "讲师",
        }
        for (r, c), value in replacements.items():
            if r < len(t0.rows) and c < len(t0.rows[r].cells):
                t0.cell(r, c).text = value

        t1 = doc.tables[1]
        en_replacements = {
            (0, 1): "Xu Chao",
            (1, 1): "Engineering",
            (2, 1): "2022213244",
            (3, 1): "Software Engineering",
            (4, 1): "School of Software",
            (5, 1): "Zheng Yongming",
            (6, 1): "Lecturer",
        }
        for (r, c), value in en_replacements.items():
            if r < len(t1.rows) and c < len(t1.rows[r].cells):
                t1.cell(r, c).text = value

    for idx, p in enumerate(doc.paragraphs[:60]):
        if "Pet Shop Management System" in p.text or "Spring Boot" in p.text:
            p.text = p.text.replace("Design and Implementation of a Pet Shop Management System Based on Spring Boot", TITLE_EN)
            p.text = p.text.replace("Pet Shop Management System Based on Spring Boot", TITLE_EN)
            p.text = p.text.replace("Spring Boot", "Node.js and JavaScript")
        if "二〇二六 年 六 月 十 日" in p.text:
            p.text = p.text.replace("二〇二六 年 六 月 十 日", "二〇二六 年 五 月 六 日")
        if "Date:" in p.text:
            p.text = "Date: 2026-05-06"
        if "签字日期： 2025年 6月 10日" in p.text:
            p.text = "                           签字日期： 2026年 5月 6日"
        if "题目" in p.text and "：" in p.text:
            p.text = p.text.split("：", 1)[0] + "：" + TITLE_CN


def remove_old_body(doc: Document) -> None:
    # The template front matter ends around this position. Removing from here keeps
    # the cover/declaration pages while replacing the abstract, TOC, and body.
    start_index = 52 if len(doc.paragraphs) > 52 else 0
    start = doc.paragraphs[start_index]._element
    parent = start.getparent()
    children = list(parent)
    idx = children.index(start)
    for child in children[idx:]:
        parent.remove(child)


def add_front_matter(doc: Document) -> None:
    add_heading(doc, "摘  要", 1)
    abstract_paras = [
        "随着移动互联网和微信生态的持续发展，小程序因其即用即走、开发成本较低、使用门槛较小等特点，逐渐成为中小型电商系统的重要载体。针对传统购物小程序在商品展示、订单管理、用户互动和客服响应方面存在的功能碎片化、体验割裂以及咨询效率不足等问题，本文设计并实现了一套基于 Node.js 与 JavaScript 的智能购物小程序系统。系统采用前后端分离的开发模式，前端基于 uni-app 开发微信小程序界面，后端基于 Express 构建 RESTful API 服务，数据库采用 MySQL 进行数据持久化存储，并结合第三方大模型能力构建 AI 智能客服模块。",
        "系统围绕“浏览商品、智能咨询、加入购物车、提交订单、订单跟踪、售后处理”的完整业务闭环展开设计，重点实现了用户登录与资料管理、商品分类展示与搜索、商品详情浏览、购物车管理、订单流转、地址维护、用户反馈、AI 聊天咨询以及后台商品、订单、售后和反馈管理等功能。与普通商品展示型小程序相比，本系统在商品详情页中引入上下文感知的智能问答能力，后端会根据商品标题、价格、规格、服务标签、发货时效等事实信息组织提示词，再调用大模型接口生成回答，并在异常情况下返回本地兜底答案，从而兼顾交互体验与服务稳定性。",
        "在系统实现过程中，本文重点讨论了前端请求封装、Pinia 状态管理、JWT 身份认证、后端分层架构、数据库事务处理、后台权限控制以及 AI 接口调用监控等关键问题。测试结果表明，系统能够较好地满足移动端购物场景下的主要功能需求，具备较清晰的模块边界、较好的可维护性和一定的实际应用价值。",
    ]
    for text in abstract_paras:
        add_para(doc, text)
    add_para(doc, "关键词：微信小程序；Node.js；Express；MySQL；AI 智能客服；电商系统")
    doc.add_page_break()

    add_heading(doc, "ABSTRACT", 1)
    english = [
        "With the continuous development of mobile Internet and the WeChat ecosystem, mini programs have become an important carrier for small and medium-sized e-commerce systems due to their lightweight usage and relatively low development cost. To address the problems of fragmented functions, inconsistent user experience, and inefficient customer service in traditional shopping mini programs, this thesis designs and implements an intelligent shopping mini-program system based on Node.js and JavaScript.",
        "The system focuses on a complete shopping workflow, including product browsing, intelligent consultation, cart management, order placement, order tracking, after-sales handling, and backend administration. The frontend is developed with uni-app and Vue-style syntax, the backend uses Express to provide RESTful APIs, and MySQL is used for persistent storage. The AI customer-service module is implemented through backend encapsulation, product-context prompt construction, fallback answers, and call logs.",
        "The implementation verifies the feasibility of a lightweight frontend-backend separated architecture in a shopping mini-program scenario. Test results show that the system can satisfy the main functional requirements and has good maintainability and practical value.",
    ]
    for text in english:
        add_para(doc, text)
    add_para(doc, "Key words: WeChat Mini Program; Node.js; Express; MySQL; AI Customer Service; E-commerce System")
    doc.add_page_break()

    add_heading(doc, "目  录", 1)
    for title, page in TOC_LINES:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = None
        if title.count(".") >= 2:
            p.paragraph_format.left_indent = Pt(24)
        elif title.count(".") == 1 and not re.match(r"^\d+\.\s", title):
            p.paragraph_format.left_indent = Pt(12)
        else:
            p.paragraph_format.left_indent = Pt(0)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        p.add_run(title)
        p.add_run("\t" + page)
    doc.add_page_break()


def clean_markdown_line(line: str) -> str:
    line = line.strip()
    line = re.sub(r"^[-*]\s+", "", line)
    line = line.replace("`", "")
    return line


def add_markdown_file(doc: Document, rel_path: str, max_paragraphs: int | None = None) -> None:
    path = DOC_DIR / rel_path
    if not path.exists():
        return
    count = 0
    in_code = False
    for raw in path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if line.startswith("```"):
            in_code = not in_code
            continue
        if in_code or not line:
            continue
        if line.startswith("# "):
            add_heading(doc, clean_markdown_line(line[2:]), 3)
        elif line.startswith("## "):
            add_heading(doc, clean_markdown_line(line[3:]), 3)
        elif "此处" in line and ("插入" in line or "预留" in line):
            add_figure_placeholder(doc, clean_markdown_line(line))
        else:
            add_para(doc, clean_markdown_line(line))
        count += 1
        if max_paragraphs and count >= max_paragraphs:
            break


def add_body(doc: Document) -> None:
    add_heading(doc, "1. 绪论", 1)
    add_heading(doc, "1.1  选题的背景", 2)
    for text in [
        "随着智能手机和移动支付的普及，消费者的购物行为从传统网页和线下门店逐步转向移动端场景。微信小程序依托微信生态，具有用户进入成本低、分享路径短、更新维护方便等特点，能够为中小型商家提供较轻量的线上经营载体。对于校园二手交易、潮流服饰、运动鞋服、生活用品等垂直购物场景而言，小程序既能满足用户随时浏览商品的需求，也能降低商家搭建独立 App 的成本。",
        "从具体业务来看，购物系统并不是单纯的商品列表展示，而是包含用户身份、商品信息、购物车、地址、订单、售后、反馈和客服咨询等多个环节。若各环节之间缺少统一的业务流转，用户在下单前后需要反复切换页面或依赖人工客服确认信息，容易降低购买效率。本课题结合实际代码实现，将用户端购物链路、后台管理链路和 AI 客服链路统一纳入系统设计，使系统更贴近真实电商业务。",
        "在大模型技术快速发展的背景下，AI 智能客服逐渐成为电商系统提升服务效率的重要方向。传统客服模块通常依赖固定问答库或关键词匹配，难以结合商品上下文进行灵活回答。本系统在商品详情页中提供 AI 咨询入口，后端根据商品事实字段生成提示词，再调用可配置的大模型供应商接口，同时保存会话消息和调用日志。该设计既能提升用户咨询体验，也为后续运营分析提供数据基础。",
    ]:
        add_para(doc, text)

    add_heading(doc, "1.2  研究现状", 2)
    for text in [
        "当前电商系统的研究与实践已经较为成熟，主流系统普遍具备商品检索、购物车、订单、支付、评价和售后等功能。大型电商平台更加强调推荐算法、供应链协同和复杂营销体系，而中小型购物小程序更关注开发成本、功能完整性和移动端交互效率。因此，在毕业设计场景下，采用轻量化技术栈构建一个功能闭环完整、结构清晰、便于扩展的购物小程序具有较强的实践意义。",
        "在前端实现方面，uni-app、Vue、Pinia 等技术能够帮助开发者以组件化方式组织移动端页面，实现跨端编译和状态共享。在后端实现方面，Node.js 与 Express 具有轻量、灵活的特点，适合构建 RESTful API，并可通过中间件机制实现统一认证、错误处理、跨域控制和安全增强。MySQL 作为关系型数据库，适合保存订单、商品、用户、售后等结构化业务数据。",
        "智能客服方面，近年来 OpenAI、DeepSeek、火山方舟等大模型接口逐渐成熟，应用系统可以通过后端服务对模型能力进行封装。相比固定问答库，大模型能够更好地理解自然语言问题，但也带来接口超时、密钥安全、回答可信度和成本监控等问题。因此，在实际系统中不能简单地把模型接口直接暴露给前端，而需要由后端进行上下文约束、异常兜底和日志记录。",
    ]:
        add_para(doc, text)

    add_heading(doc, "1.3  研究意义", 2)
    for text in [
        "本课题的实践意义在于完整呈现一个购物小程序从前端页面到后端接口、再到数据库和 AI 能力接入的实现过程。系统不是停留在静态页面展示层面，而是实现了用户登录、商品查询、购物车同步、订单创建、地址选择、售后申请、反馈提交、后台管理和 AI 咨询等真实业务功能，能够体现软件工程专业在需求分析、系统设计、编码实现和测试验证方面的综合训练。",
        "从工程结构角度看，本系统采用前后端分离和业务模块分层的方式组织代码。前端页面通过 services 层调用后端接口，后端按照 routes、controller、service、repository 的层次拆分业务，能够降低页面逻辑、业务规则和数据库访问之间的耦合度。该结构有利于后续扩展支付、优惠券、推荐、消息通知等模块，也便于定位和维护问题。",
        "从应用价值角度看，AI 客服模块使用户能够围绕商品规格、材质、发货、售后等问题直接提问，减少人工客服压力。后台管理端提供商品维护、订单推进、售后处理、反馈查看和运营概览等功能，使商家能够对商品和交易过程进行集中管理。系统整体对中小型购物场景具有一定参考价值。",
    ]:
        add_para(doc, text)

    add_heading(doc, "1.4  论文结构", 2)
    add_para(doc, "本文共分为六章。第一章介绍选题背景、研究现状和研究意义；第二章介绍系统开发工具、开发技术和运行环境；第三章对系统进行可行性分析、功能需求分析、性能需求分析和安全需求分析；第四章阐述系统总体功能、架构、业务流程和数据库设计；第五章结合前后端代码说明用户端、管理端以及后端核心模块的具体实现；第六章介绍系统测试目标、测试方法、测试用例和测试结果。最后给出结论、致谢和参考文献。")

    add_heading(doc, "2. 系统开发技术与环境", 1)
    add_heading(doc, "2.1  开发工具", 2)
    for text in [
        "本系统主要在 Windows 本地环境下开发，使用 Visual Studio Code 进行前端页面和后端服务代码编写，使用微信小程序开发工具或 H5 调试环境进行页面预览，使用 MySQL 数据库工具进行数据表维护和数据检查。项目代码分为 frontend/miniprogram 和 backend 两部分，分别负责小程序端和接口服务端。",
        "在项目调试过程中，前端通过 npm 脚本启动 uni-app 调试服务，后端通过 Node.js 启动 Express API 服务，数据库通过迁移脚本和种子数据完成初始化。该开发方式便于分别调试页面交互、接口响应和数据库数据，也便于在论文中按照模块说明系统实现。",
    ]:
        add_para(doc, text)

    add_heading(doc, "2.2  开发技术", 2)
    tech_sections = [
        ("2.2.1  uni-app 小程序开发技术", "前端采用 uni-app 构建微信小程序页面。项目页面集中在 frontend/miniprogram/src/pages 目录，包括登录、首页、商品详情、购物车、订单确认、订单列表、地址、售后、反馈、聊天和后台管理等页面。uni-app 允许使用 Vue 风格语法组织页面逻辑，使模板、脚本和样式能够在同一个 .vue 文件中维护，适合本系统这种页面较多、交互链路较完整的移动端应用。"),
        ("2.2.2  Vue 与 Pinia 状态管理", "前端使用 Pinia 管理用户会话和购物车状态。user store 从本地缓存读取 shopping_access_token 和 shopping_user_profile，提供 isLoggedIn、isAdmin 等状态判断，并在登录、资料更新和退出登录时同步本地缓存。cart store 则维护购物车列表、商品总数和金额，在添加、修改数量和删除商品后立即用后端返回结果刷新状态，避免页面显示与数据库数据不一致。"),
        ("2.2.3  Node.js 与 Express 后端技术", "后端采用 Node.js 与 Express 构建 RESTful API。系统按 auth、product、cart、order、chat、address、afterSale、feedback、admin 等业务域拆分模块，每个模块通常包含 routes、controller、service、repository 文件。routes 负责绑定路由和中间件，controller 负责接收参数并返回响应，service 负责处理业务规则，repository 负责数据库读写。"),
        ("2.2.4  MySQL 数据库技术", "系统使用 MySQL 保存结构化业务数据。商品、商品图集、用户、购物车、地址、订单、订单明细、订单物流、订单时间线、售后申请、用户反馈、聊天会话、聊天消息以及 AI 调用日志都保存在数据库中。订单创建和后台商品维护等场景使用事务，保证主表、明细表和关联数据的一致性。"),
        ("2.2.5  AI 智能客服接口技术", "AI 客服能力由后端 integrations/ai/aiClient.js 统一封装。系统支持根据配置选择 deepseek、openai 或 ark 等供应商，调用时设置模型、temperature、system prompt 和用户问题，并通过 AbortSignal.timeout 控制超时。若接口异常、密钥缺失或返回空结果，后端会生成本地兜底答案，并将 provider、model、success、fallback_used、latency_ms 和 error_message 写入 ai_call_logs。"),
    ]
    for title, text in tech_sections:
        add_heading(doc, title, 3)
        add_para(doc, text)
    add_heading(doc, "2.3  系统运行环境", 2)
    add_para(doc, "系统运行环境包括 Node.js、npm、MySQL、微信小程序调试环境和浏览器 H5 调试环境。后端服务默认提供 /api 前缀的接口，前端请求层通过 VITE_API_BASE_URL 或默认地址 http://localhost:3001/api 访问后端。前端 request.ts 会自动从本地缓存读取 token 并附加 Authorization: Bearer token 请求头，非 GET 请求统一设置 JSON 请求头，后端返回 HTTP 400 以上状态时前端会将错误信息转换为异常。")

    add_heading(doc, "3. 系统分析", 1)
    add_heading(doc, "3.1  可行性分析", 2)
    for title, text in [
        ("3.1.1  技术可行性", "本系统采用的 uni-app、Vue、Pinia、Node.js、Express 和 MySQL 都是成熟技术，社区资料较多，能够满足毕业设计中购物小程序的前端展示、接口服务和数据存储需求。AI 客服模块通过后端代理调用大模型接口，不要求前端直接接触密钥，能够在现有架构上平滑接入。"),
        ("3.1.2  经济可行性", "系统主要使用开源框架和本地开发环境完成，数据库、后端服务和前端调试均可在个人电脑上运行。除第三方 AI 接口可能产生调用费用外，系统开发和测试成本较低，符合毕业设计项目的资源约束。"),
        ("3.1.3  运行可行性", "系统面向普通购物用户和后台管理员两类角色，用户端功能围绕浏览、加购、下单、售后和咨询展开，管理端功能围绕商品、订单、售后、反馈和统计展开。页面入口清晰，业务流程符合常见购物习惯，具备较好的运行可行性。"),
    ]:
        add_heading(doc, title, 3)
        add_para(doc, text)

    add_heading(doc, "3.2  需求分析", 2)
    add_heading(doc, "3.2.1  功能需求分析", 3)
    for text in [
        "普通用户需要完成登录注册、商品浏览、商品搜索与筛选、商品详情查看、规格和数量选择、加入购物车、购物车数量修改、订单确认、地址管理、订单查看、售后申请、反馈提交、AI 智能客服咨询和个人资料维护等操作。商品详情页需要展示商品主图、图集、标题、副标题、价格、原价、库存、销量、规格、详情说明和用户评价，并提供 AI 咨询、加入购物车和立即购买入口。",
        "管理员需要完成后台登录、商品新增与编辑、商品上下架、库存维护、订单状态更新、物流信息录入、售后状态处理、用户反馈查看与状态维护、运营指标查看等操作。后台概览需要展示日订单、待处理售后、在售商品、待看反馈、热门商品、品类销量、售后原因排行以及 AI 调用监控等数据。",
        "后端功能需求包括 JWT 登录认证、普通用户接口保护、管理员接口保护、商品检索、购物车同步、订单事务创建、地址增删改查、售后申请、反馈保存、聊天会话和消息保存、AI 接口调用、异常兜底、统一错误响应以及请求 ID 追踪。数据库需要支持多表关联查询和事务提交，保证订单、购物车和售后数据之间的业务一致性。",
    ]:
        add_para(doc, text)
    add_figure_placeholder(doc, "图3-1 用户端功能结构图")
    add_figure_placeholder(doc, "图3-2 管理端功能结构图")

    add_heading(doc, "3.2.2  性能需求分析", 3)
    add_para(doc, "系统高频接口包括商品列表、商品搜索、商品详情、购物车、订单列表和后台概览。商品检索接口需要支持分页，product.repository.js 将 pageSize 限制在 1 到 100 之间，并通过 LIMIT 和 OFFSET 控制返回范围。后台概览通过 Promise.all 并发读取商品、订单、售后、趋势和 AI 监控数据，减少串行等待时间。AI 接口由于受外部服务影响较大，需要设置超时和降级逻辑，避免长时间阻塞用户聊天页面。")
    add_heading(doc, "3.2.3  安全需求分析", 3)
    add_para(doc, "系统涉及用户资料、收货地址、订单、售后、反馈和聊天记录等私有数据，必须进行身份认证和权限控制。后端 authMiddleware.js 提供 optionalAuth、requireAuth 和 requireAdmin 三类中间件，普通用户私有接口必须通过 requireAuth 校验，后台接口在 requireAuth 后还需要 requireAdmin 校验角色。JWT 载荷包含 userId、role 和 nickname，有效期设置为 7 天。AI 接口密钥只保存在后端环境变量中，前端不能直接获取。")
    add_heading(doc, "3.2.4  兼容性需求分析", 3)
    add_para(doc, "系统前端需要适配微信小程序和 H5 调试环境。页面布局采用 rpx、flex、滚动容器和组件化结构，购物车、详情页、聊天页、后台列表页等主要页面都围绕移动端窄屏场景设计。接口层通过统一 request.ts 封装基础地址、token 和错误处理，降低不同环境下切换接口地址的成本。")

    add_heading(doc, "4. 系统设计", 1)
    add_heading(doc, "4.1  系统总体功能设计", 2)
    add_para(doc, "系统总体功能分为用户端、管理端和后端服务三部分。用户端提供购物链路和 AI 咨询体验，管理端提供运营维护能力，后端服务负责认证、业务规则、数据库访问和第三方 AI 能力接入。三部分通过 RESTful API 连接，前端不直接访问数据库，也不直接调用 AI 服务。")
    add_figure_placeholder(doc, "图4-1 系统总体功能模块图")
    add_heading(doc, "4.2  系统架构设计", 2)
    add_para(doc, "系统采用前后端分离架构。前端页面层依赖组件层、状态层和 services 层；services 层统一调用 request.ts；request.ts 负责拼接 API 地址、携带 token、设置请求头和解析响应。后端 API 层通过 routes 接收请求，再进入 controller、service 和 repository。repository 通过 config/db.js 中的 query 和 withTransaction 与 MySQL 交互。AI 能力作为外部服务被封装在 integrations/ai/aiClient.js 中，业务模块只调用 askAi 方法。")
    add_figure_placeholder(doc, "图4-2 系统总体架构图")
    add_heading(doc, "4.3  系统业务流程设计", 2)
    for text in [
        "用户下单流程为：用户在首页或搜索页选择商品，进入详情页查看图集、规格、评价和详情说明；选择规格和数量后调用购物车接口加入购物车；进入购物车页面确认商品数量和金额；进入订单确认页读取默认地址和购物车结算项；提交订单时后端在事务中创建 orders、order_items 和 order_timeline 记录，并清空购物车。",
        "AI 咨询流程为：用户在商品详情页点击 AI 入口，前端携带 productId 创建聊天会话；后端从数据库读取商品上下文字段并保存 chat_sessions；用户发送问题后，后端先保存用户消息，再构建商品提示词调用模型接口；获得回答后保存 assistant 消息并返回前端。如果模型接口失败，后端返回由商品规格、服务标签、发货时效等字段组成的兜底回答，同时记录调用日志。",
        "后台订单处理流程为：管理员进入订单列表，查看订单状态、金额、商品和创建时间；进入订单详情后可更新订单状态或录入物流信息。后端 updateOrderStatus 会在事务中更新 orders.status，并根据状态生成 order_timeline 记录；updateOrderLogistics 会写入或更新 order_shipments，并追加物流时间线，保证后台处理行为可追踪。",
    ]:
        add_para(doc, text)
    add_figure_placeholder(doc, "图4-3 下单业务流程图")
    add_figure_placeholder(doc, "图4-4 AI 咨询业务流程图")

    add_heading(doc, "4.4  系统数据库设计", 2)
    add_heading(doc, "4.4.1  数据库需求分析", 3)
    add_para(doc, "数据库需要支撑用户身份、商品展示、购物车结算、订单状态流转、地址管理、售后申请、反馈管理、聊天记录和 AI 调用监控。不同表之间通过用户 ID、商品 ID、订单 ID 和会话 ID 建立关系，既满足业务查询，也保证数据之间的可追溯性。")
    add_heading(doc, "4.4.2  数据库概念结构设计", 3)
    add_para(doc, "用户与地址、购物车、订单、售后、反馈和聊天会话是一对多关系；商品与商品图集、购物车条目、订单明细、反馈和 AI 咨询上下文存在关联；订单与订单明细、订单时间线、物流信息和售后申请存在一对多或一对一关系；AI 调用日志与聊天模块相对独立，用于后台监控供应商调用成功率、兜底次数和平均延迟。")
    add_figure_placeholder(doc, "图4-5 数据库 E-R 图")
    add_heading(doc, "4.4.3  数据表设计", 3)
    add_table(
        doc,
        ["数据表", "主要字段或内容", "作用"],
        [
            ["users", "id、nickname、phone、avatar、role", "保存用户身份、角色和资料"],
            ["products", "title、price、stock、sales、category、badges_json、sizes_json", "保存商品基础信息和营销属性"],
            ["product_galleries", "product_id、image_url、sort_no", "保存商品图集"],
            ["cart_items", "user_id、product_id、quantity、size", "保存购物车条目"],
            ["addresses", "user_id、name、phone、region、detail、is_default", "保存收货地址"],
            ["orders", "user_id、status、amount、address_id、remark", "保存订单主信息"],
            ["order_items", "order_id、product_id、quantity、product_price", "保存订单商品快照"],
            ["order_timeline", "order_id、title、description", "保存订单状态流转记录"],
            ["after_sales", "order_id、product_title、reason、status", "保存售后申请"],
            ["feedbacks", "user_id、product_id、summary、status", "保存用户反馈与评价"],
            ["chat_sessions", "user_id、product_id", "保存 AI 会话"],
            ["chat_messages", "session_id、role、content", "保存聊天消息"],
            ["ai_call_logs", "provider、model、success、fallback_used、latency_ms", "保存 AI 调用监控数据"],
        ],
        "表4-1 系统主要数据表",
    )

    add_heading(doc, "5. 系统功能实现", 1)
    add_heading(doc, "5.1  用户端功能的实现", 2)
    feature_paras = [
        ("5.1.1  用户登录注册", "登录功能由 frontend/miniprogram/src/pages/auth/login.vue、authService.ts、user store 和后端 auth 模块共同完成。前端调用 loginWithIdentity 提交 code、nickname 和 identity，后端 auth.service.js 根据 identity 查找演示用户，构造包含 userId、role、nickname 的 JWT 并返回 accessToken。user store 将 token 和用户资料写入本地缓存，后续 request.ts 会自动读取 token 并附加到 Authorization 请求头。"),
        ("5.1.2  首页商品展示与搜索", "商品列表接口由 product.repository.js 实现，支持 keyword、scene、category、onSaleOnly、sort、page 和 pageSize 等查询条件。keyword 会同时匹配 title、subtitle 和 category；scene 为 newIn48h 时筛选 48 小时上新商品，scene 为 buyerFavorite 时筛选高好评或高销量商品；sort 支持价格升序、价格降序、热度和综合排序。列表返回前还会根据商品 ID 批量读取 product_galleries，避免前端重复请求图集。"),
        ("5.1.3  商品详情与反馈", "商品详情页 detail.vue 在 onLoad 时读取商品 ID，并通过 Promise.all 同时请求商品详情和用户反馈。页面展示主图、缩略图、价格、库存、规格、详情说明、销量、评价列表和 AI 咨询入口。用户未登录时点击写评价、加购或 AI 咨询会跳转登录页，已登录用户可以进入反馈页面提交评价，也可以进入聊天页面围绕当前商品提问。"),
        ("5.1.4  购物车管理", "购物车页面 cart/index.vue 在 onMounted 和 onShow 时加载购物车，未登录时显示 EmptyStateCard 引导登录，已登录时展示商品卡片、规格、价格、数量步进器和移除按钮。cart store 的 fetchCart、add、changeQuantity 和 remove 都会调用后端接口，并使用返回的 list、totalCount 和 totalAmount 刷新本地状态。数量修改时前端通过 Math.max(1, quantity) 保证数量不小于 1。"),
        ("5.1.5  订单确认与地址管理", "订单确认页调用 orderRepository.getPreview 获取默认地址、购物车结算项和金额汇总。后端 buildSummary 根据商品数量与价格计算 goodsAmount，并设置运费和优惠金额，最终得到 payableAmount。用户提交订单时，后端读取指定地址或默认地址，在事务中写入订单主表、订单明细和订单时间线，并清空购物车，保证下单链路数据一致。地址模块提供列表、新增、更新和删除能力，若更新或删除不存在地址，service 层会抛出 ADDRESS_NOT_FOUND。"),
        ("5.1.6  售后申请", "售后模块面向已登录用户开放。用户在订单相关页面发起售后申请时，后端 afterSale.service.js 调用 repository 创建申请记录；如果订单不存在或不属于当前用户，则返回 ORDER_NOT_FOUND。后台管理员可以查看售后列表并更新售后状态，系统概览中也会统计待处理售后数量和售后原因排行，便于运营人员优先处理异常订单。"),
        ("5.1.7  AI 智能客服聊天", "聊天页面通过 chatService 创建会话并发送问题。后端 chat.service.js 首先校验会话是否属于当前用户，再保存用户消息；若会话绑定商品，则从数据库读取商品上下文，包括标题、副标题、分类、价格、详情、标签、规格、服务保障、好评率和发货时效。aiClient.js 将这些事实写入系统提示词，要求模型在商品事实范围内回答，并对无关问题进行引导。模型调用成功后保存 assistant 消息，失败时返回本地兜底答案并写入 ai_call_logs。"),
        ("5.1.8  个人中心与资料管理", "个人中心依赖 user store 判断登录状态和角色。普通用户可以查看订单、地址、售后、反馈和资料维护入口；管理员用户可以看到后台管理入口。用户更新昵称、头像或角色时，前端调用 updateAuthProfile，后端更新资料后重新签发 token，前端再同步本地缓存，从而保证页面展示、接口权限和本地会话保持一致。"),
    ]
    for title, text in feature_paras:
        add_heading(doc, title, 3)
        add_para(doc, text)
        if "商品详情" in title:
            add_figure_placeholder(doc, "图5-1 商品详情页预留截图")
        if "购物车" in title:
            add_figure_placeholder(doc, "图5-2 购物车页面预留截图")
        if "AI" in title:
            add_figure_placeholder(doc, "图5-3 AI 智能客服页面预留截图")

    add_heading(doc, "5.2  管理端功能实现", 2)
    admin_paras = [
        ("5.2.1  后台登录与权限控制", "后台功能只能由管理员角色访问。前端 user store 通过 profile.role 判断 isAdmin，页面入口据此显示；后端 admin 路由必须经过 requireAuth 和 requireAdmin，若 token 缺失返回 401，若角色不是 admin 返回 403。该设计避免普通用户通过直接访问后台接口修改商品或订单。"),
        ("5.2.2  商品管理", "后台商品管理由 admin.repository.db.js 实现数据库读写。新增商品时系统生成 productId，将标题、副标题、价格、原价、库存、分类、封面、详情、上下架状态、标签、规格、服务保障、上新标记、好评率和发货时效写入 products 表，并在同一事务中写入 product_galleries。更新商品时，如果修改了图集或封面，系统会先删除旧图集再按顺序写入新图集，保证详情页展示数据与后台维护结果一致。"),
        ("5.2.3  订单管理", "后台订单管理支持列表查看、详情查看、状态更新和物流录入。订单列表会聚合订单主信息与订单明细，详情页额外展示订单时间线和物流信息。管理员更新状态时，后端根据 pending_payment、pending_shipping、shipped、completed、cancelled 等状态生成对应时间线文案；录入物流时使用 ON DUPLICATE KEY UPDATE 保证同一订单物流记录可重复更新。"),
        ("5.2.4  售后与反馈管理", "售后管理用于查看用户提交的售后原因、订单号、商品名称、提交时间和处理状态，管理员可以更新状态。反馈管理用于查看商品评价或体验建议，并可维护反馈状态。系统概览中会统计待处理售后和待看反馈，帮助管理员识别需要优先处理的事项。"),
        ("5.2.5  系统数据统计与概览", "后台概览通过 adminRepository.getOverview 聚合多类运营数据，包括今日订单数、待处理售后、在售商品数、待看反馈、热门商品、品类销量排行、售后原因排行、订单趋势和 AI Agent 调用趋势。AI 监控数据来自 ai_call_logs，可计算今日调用次数、兜底次数、成功率和平均延迟，为判断智能客服稳定性提供依据。"),
    ]
    for title, text in admin_paras:
        add_heading(doc, title, 3)
        add_para(doc, text)
    add_figure_placeholder(doc, "图5-4 管理端概览页面预留截图")

    add_heading(doc, "5.3  后端核心模块实现", 2)
    add_para(doc, "为了进一步保证论文内容与代码一致，本节结合项目中的模块文档和源码实现，对后端核心模块进行补充说明。以下内容对应 backend/src/modules 下的 auth、product、cart、order、address、afterSale、feedback、chat 和 admin 等模块，以及 frontend/miniprogram/src/services、src/store 和 src/pages 中的前端调用关系。")
    for rel in [
        "02-用户认证与个人中心模块.md",
        "03-首页商品展示与搜索模块.md",
        "04-商品详情与反馈模块.md",
        "05-购物车模块.md",
        "06-订单与地址模块.md",
        "07-售后申请模块.md",
        "08-AI智能客服模块.md",
        "09-后台管理模块.md",
        "10-系统安全与接口规范.md",
    ]:
        add_markdown_file(doc, rel)

    add_heading(doc, "6. 系统测试", 1)
    add_heading(doc, "6.1  测试目标与范围", 2)
    add_para(doc, "系统测试的目标是验证购物小程序在主要业务链路上的功能正确性、接口稳定性、权限控制有效性和页面兼容性。测试范围包括用户登录、商品列表、商品搜索、商品详情、购物车、订单确认、地址管理、售后申请、用户反馈、AI 聊天、后台商品管理、后台订单管理、售后反馈处理和后台概览统计。")
    add_heading(doc, "6.2  测试环境与工具", 2)
    add_para(doc, "测试环境与开发环境一致，主要包括 Windows、本地 Node.js 服务、MySQL 数据库、uni-app 调试环境和浏览器或微信开发者工具。接口测试可结合浏览器控制台、前端页面操作和后端日志进行观察；数据库测试通过查看 orders、order_items、cart_items、chat_messages、ai_call_logs 等表确认数据是否正确写入。")
    add_heading(doc, "6.3  测试方法与策略", 2)
    add_para(doc, "测试采用黑盒功能测试与代码路径检查相结合的方式。用户端以真实购物流程为主线，从登录、浏览、详情、加购、结算、下单、售后和 AI 咨询依次验证；后台端以管理员处理流程为主线，验证商品维护、订单状态流转、物流录入、售后反馈维护和统计展示。对于订单创建、AI 调用和权限控制等风险较高的功能，还需要结合数据库和后端代码检查事务、兜底和鉴权逻辑是否符合预期。")
    add_heading(doc, "6.4  功能测试", 2)
    add_heading(doc, "6.4.1  核心功能测试用例表", 3)
    add_table(
        doc,
        ["编号", "测试功能", "测试步骤", "预期结果"],
        [
            ["TC01", "用户登录", "选择普通用户身份登录", "返回 accessToken，个人中心显示用户信息"],
            ["TC02", "商品搜索", "输入关键词并切换排序", "列表按关键词和排序条件返回商品"],
            ["TC03", "商品详情", "进入详情页并切换缩略图", "展示商品信息、图集、规格和评价"],
            ["TC04", "购物车", "加入商品、修改数量、移除商品", "购物车数量和金额同步更新"],
            ["TC05", "订单创建", "选择地址并提交订单", "生成订单、订单明细和时间线，购物车清空"],
            ["TC06", "AI 咨询", "在商品详情页发起提问", "返回商品相关回答并保存聊天记录"],
            ["TC07", "后台商品管理", "新增或编辑商品", "商品表和图集表数据同步更新"],
            ["TC08", "后台订单状态", "更新订单状态并录入物流", "订单状态、物流和时间线正确写入"],
            ["TC09", "权限控制", "普通用户访问后台接口", "后端返回 403 或拒绝访问"],
        ],
        "表6-1 核心功能测试用例表",
    )
    add_heading(doc, "6.4.2  测试结论分析", 3)
    add_para(doc, "从功能测试结果看，系统主要业务流程能够按预期运行。商品模块能够根据关键词、分类、场景和排序条件返回数据；购物车模块在添加、修改数量和删除后能够同步金额；订单模块能够在事务中生成订单、明细和时间线并清空购物车；AI 聊天模块能够保存消息并在模型异常时返回兜底内容；后台模块能够完成商品、订单、售后、反馈和概览数据管理。")
    add_heading(doc, "6.5  非功能测试", 2)
    for title, text in [
        ("6.5.1  兼容性测试", "在 H5 调试环境和微信小程序调试环境中检查主要页面，重点观察首页、详情页、购物车、订单确认页、聊天页和后台管理页的布局是否正常，滚动区域和底部操作栏是否遮挡内容。"),
        ("6.5.2  响应时效与接口稳定性测试", "商品列表和详情接口返回数据较快；后台概览通过并发查询降低等待时间；AI 接口由于依赖第三方服务，系统通过 timeout 和 fallback 机制保证用户不会长时间等待。"),
        ("6.5.3  安全机制与权限稳定性测试", "未登录用户访问购物车、订单、地址、售后和聊天接口时应被拒绝；普通用户访问后台接口时应返回无权限；管理员 token 能够访问后台管理接口。"),
        ("6.5.4  运行稳定性观察", "连续执行登录、加购、下单、售后、反馈和 AI 聊天操作后，数据库中的订单、购物车、售后、反馈、聊天消息和 AI 日志能够保持一致，没有出现明显的数据丢失或重复写入问题。"),
    ]:
        add_heading(doc, title, 3)
        add_para(doc, text)
    add_heading(doc, "6.6  测试结果分析与总结", 2)
    add_para(doc, "综合功能测试和非功能测试可知，本系统已经实现购物小程序的核心业务闭环，前端页面、后端接口和数据库之间的数据流转基本正确。系统仍可进一步完善真实支付、消息通知、库存并发扣减和更精细的 AI 内容审核等功能，但作为毕业设计系统，已经能够较完整地展示移动端购物系统的需求分析、系统设计、编码实现和测试验证过程。")
    add_markdown_file(doc, "11-系统测试与运行部署.md")

    add_heading(doc, "结论", 1)
    add_para(doc, "本文围绕基于 Node.js 与 JavaScript 的购物小程序系统展开研究与实现，完成了从需求分析、系统设计、数据库设计到前后端功能实现和系统测试的全过程。系统前端采用 uni-app 和 Vue 风格语法构建小程序页面，后端采用 Express 提供 RESTful API，数据库采用 MySQL 保存业务数据，并通过 AI 接口封装实现智能客服能力。")
    add_para(doc, "从实现结果看，系统完成了普通用户端的登录、商品浏览、搜索筛选、商品详情、购物车、订单、地址、售后、反馈、AI 咨询和个人中心等功能，也完成了管理端的商品管理、订单管理、售后管理、反馈管理和运营概览等功能。后端通过 JWT、requireAuth、requireAdmin、事务处理、统一请求封装和 AI 调用日志等机制，提高了系统的安全性、一致性和可维护性。")
    add_para(doc, "当然，系统仍存在可以继续完善的方向。例如，当前支付流程以订单创建和状态推进为主，尚未接入真实支付网关；库存并发控制可以进一步结合行级锁或库存流水；AI 客服可以加入更精细的内容审核、知识库召回和人工客服转接能力。后续若继续扩展，可在现有分层架构基础上逐步加入这些能力。")

    add_heading(doc, "致  谢", 1)
    add_para(doc, "在本论文和毕业设计系统完成过程中，我得到了指导老师郑勇明讲师的耐心指导。老师在选题方向、系统结构、论文撰写和功能完善等方面给予了许多建议，使我能够更加系统地理解软件工程项目从需求到实现的完整过程。在此向老师表示衷心感谢。")
    add_para(doc, "同时感谢学院各位老师在大学期间的培养，使我掌握了前端开发、后端开发、数据库设计和软件工程实践等知识。感谢同学和朋友在系统测试、问题反馈和论文修改过程中的帮助。最后感谢家人在学习和生活上的支持，使我能够顺利完成本次毕业设计。")

    add_heading(doc, "参考文献", 1)
    refs = [
        "[1] 尤雨溪. Vue.js 官方文档[EB/OL].",
        "[2] DCloud. uni-app 官方文档[EB/OL].",
        "[3] Express.js. Express Web Framework Documentation[EB/OL].",
        "[4] MySQL. MySQL Reference Manual[EB/OL].",
        "[5] OpenJS Foundation. Node.js Documentation[EB/OL].",
        "[6] JSON Web Token. JWT Introduction[EB/OL].",
        "[7] Pinia. Pinia Official Documentation[EB/OL].",
        "[8] OpenAI. API Documentation[EB/OL].",
        "[9] DeepSeek. DeepSeek API Documentation[EB/OL].",
        "[10] 软件工程导论相关教材与课程资料.",
    ]
    for ref in refs:
        p = add_para(doc, ref)
        p.paragraph_format.first_line_indent = None


def main() -> None:
    source = BACKUP if BACKUP.exists() else FALLBACK_TEMPLATE
    doc = Document(str(source))
    set_doc_defaults(doc)
    update_cover(doc)
    remove_old_body(doc)
    add_front_matter(doc)
    add_body(doc)
    doc.save(str(OUTPUT))
    try:
        doc.save(str(DESKTOP))
    except PermissionError:
        doc.save(str(DESKTOP_ALT))
        print(f"desktop_fallback={DESKTOP_ALT}")
    print(f"saved={OUTPUT}")
    print(f"desktop={DESKTOP}")


if __name__ == "__main__":
    main()
