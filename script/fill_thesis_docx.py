# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document


ROOT = Path(r"C:\Users\86157\Desktop\UniAppPencil")
DOC_DIR = ROOT / "毕设文档"
TEMPLATE = DOC_DIR / "毕业论文模板.docx"
OUTPUT = DOC_DIR / "毕业论文_徐超_初稿.docx"
ROOT_DRAFT = ROOT / "论文模板.md"

TITLE_CN = "基于 Node.js 与 JavaScript 实现购物小程序的设计与实现"
TITLE_EN = "Design and Implementation of a Shopping Mini-Program Based on Node.js and JavaScript"


def markdown_to_blocks(text: str) -> list[tuple[int, str]]:
    blocks: list[tuple[int, str]] = []
    in_code = False
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            continue
        level = 0
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            line = line.lstrip("#").strip()
        line = line.replace("**", "").replace("`", "")
        if line:
            blocks.append((level, line))
    return blocks


def build_body() -> list[tuple[int, str]]:
    parts: list[tuple[int, str]] = []

    if ROOT_DRAFT.exists():
        draft = ROOT_DRAFT.read_text(encoding="utf-8")
        marker = "## 5 系统测试"
        if marker in draft:
            before, after = draft.split(marker, 1)
            parts.extend(markdown_to_blocks(before))
        else:
            after = ""
            parts.extend(markdown_to_blocks(draft))

    module_files = [
        "01-系统总体架构与技术选型.md",
        "02-用户认证与个人中心模块.md",
        "03-首页商品展示与搜索模块.md",
        "04-商品详情与反馈模块.md",
        "05-购物车模块.md",
        "06-订单与地址模块.md",
        "07-售后申请模块.md",
        "08-AI智能客服模块.md",
        "09-后台管理模块.md",
        "10-系统安全与接口规范.md",
        "11-系统测试与运行部署.md",
    ]

    parts.append((3, "4.7 系统各模块详细实现说明"))
    parts.append(
        (
            0,
            "本节根据项目代码结构和毕设文档中的模块说明，对系统各模块实现细节进行补充说明，便于后续在 Word 中继续扩写和插入运行截图。",
        )
    )

    for name in module_files:
        path = DOC_DIR / name
        if path.exists():
            parts.extend(markdown_to_blocks(path.read_text(encoding="utf-8")))

    if ROOT_DRAFT.exists() and "## 5 系统测试" in ROOT_DRAFT.read_text(encoding="utf-8"):
        parts.extend(markdown_to_blocks("## 5 系统测试" + after))

    return parts


def fill_table(table, rows: list[list[str]]) -> None:
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.text = ""
    for r_idx, data_row in enumerate(rows[: len(table.rows)]):
        for c_idx, value in enumerate(data_row[: len(table.columns)]):
            table.cell(r_idx, c_idx).text = value


def simple_schema_table(title: str, fields: list[tuple[str, str, str, str]]) -> list[list[str]]:
    rows = [["字段", "数据类型", "字段说明", "备注"]]
    rows.extend([list(item) for item in fields])
    return rows


def update_tables(doc: Document) -> None:
    if len(doc.tables) >= 2:
        t0 = doc.tables[0]
        t0.cell(0, 1).text = TITLE_CN
        t0.cell(1, 1).text = TITLE_EN
        t0.cell(2, 1).text = "徐超"
        t0.cell(2, 3).text = "工学"
        t0.cell(3, 1).text = "2022213244"
        t0.cell(4, 1).text = "软件工程"
        t0.cell(5, 1).text = "软件学院"
        t0.cell(6, 1).text = "郑勇明"
        t0.cell(6, 3).text = "讲师"

        t1 = doc.tables[1]
        t1.cell(0, 1).text = "Xu Chao"
        t1.cell(1, 1).text = "Engineering"
        t1.cell(2, 1).text = "2022213244"
        t1.cell(3, 1).text = "Software Engineering"
        t1.cell(4, 1).text = "School of Software"
        t1.cell(5, 1).text = "Zheng Yongming"
        t1.cell(6, 1).text = "Lecturer"

    schema_tables = [
        simple_schema_table("users", [
            ("id", "varchar(40)", "用户主键", "非空，唯一"),
            ("nickname", "varchar(50)", "用户昵称", "可空"),
            ("avatar", "varchar(255)", "头像地址", "可空"),
            ("role", "varchar(20)", "角色标识", "user/admin"),
            ("phone", "varchar(20)", "手机号", "可空"),
            ("status", "tinyint", "账号状态", "默认启用"),
            ("created_at", "datetime", "创建时间", "非空"),
            ("updated_at", "datetime", "更新时间", "非空"),
        ]),
        simple_schema_table("products", [
            ("id", "varchar(40)", "商品主键", "非空，唯一"),
            ("title", "varchar(100)", "商品标题", "非空"),
            ("category", "varchar(50)", "商品分类", "用于筛选"),
            ("price", "decimal(10,2)", "销售价格", "非空"),
            ("original_price", "decimal(10,2)", "原价", "可空"),
            ("stock", "int", "库存数量", "非空"),
            ("cover", "varchar(255)", "商品封面图", "可空"),
            ("badges_json", "json", "商品标签", "新品/热卖等"),
        ]),
        simple_schema_table("product_galleries", [
            ("id", "varchar(40)", "图集主键", "非空"),
            ("product_id", "varchar(40)", "商品ID", "关联 products"),
            ("image_url", "varchar(255)", "图片地址", "非空"),
            ("sort_order", "int", "排序值", "越小越靠前"),
            ("created_at", "datetime", "创建时间", "非空"),
        ]),
        simple_schema_table("cart_items", [
            ("id", "varchar(40)", "购物车条目ID", "非空"),
            ("user_id", "varchar(40)", "用户ID", "关联 users"),
            ("product_id", "varchar(40)", "商品ID", "关联 products"),
            ("size", "varchar(50)", "规格", "可空"),
            ("quantity", "int", "购买数量", "非空"),
            ("created_at", "datetime", "创建时间", "非空"),
        ]),
        simple_schema_table("addresses", [
            ("id", "varchar(40)", "地址ID", "非空"),
            ("user_id", "varchar(40)", "用户ID", "关联 users"),
            ("receiver", "varchar(50)", "收货人", "非空"),
            ("phone", "varchar(20)", "联系电话", "非空"),
            ("region", "varchar(100)", "省市区", "非空"),
            ("detail", "varchar(255)", "详细地址", "非空"),
            ("is_default", "tinyint", "默认地址", "0/1"),
        ]),
        simple_schema_table("orders", [
            ("id", "varchar(40)", "订单ID", "非空"),
            ("order_no", "varchar(64)", "订单编号", "唯一"),
            ("user_id", "varchar(40)", "用户ID", "关联 users"),
            ("address_snapshot_json", "json", "地址快照", "下单时保存"),
            ("total_amount", "decimal(10,2)", "订单金额", "非空"),
            ("status", "varchar(20)", "订单状态", "待付款/已发货等"),
            ("remark", "varchar(255)", "订单备注", "可空"),
        ]),
        simple_schema_table("order_items", [
            ("id", "varchar(40)", "明细ID", "非空"),
            ("order_id", "varchar(40)", "订单ID", "关联 orders"),
            ("product_id", "varchar(40)", "商品ID", "关联 products"),
            ("product_title", "varchar(100)", "商品名称快照", "非空"),
            ("quantity", "int", "购买数量", "非空"),
            ("unit_price", "decimal(10,2)", "成交单价", "非空"),
        ]),
        simple_schema_table("after_sales", [
            ("id", "varchar(40)", "售后ID", "非空"),
            ("user_id", "varchar(40)", "用户ID", "关联 users"),
            ("order_id", "varchar(40)", "订单ID", "关联 orders"),
            ("reason", "varchar(100)", "售后原因", "非空"),
            ("description", "text", "问题说明", "可空"),
            ("status", "varchar(20)", "处理状态", "待处理/已处理"),
        ]),
        simple_schema_table("feedbacks", [
            ("id", "varchar(40)", "反馈ID", "非空"),
            ("user_id", "varchar(40)", "用户ID", "关联 users"),
            ("product_id", "varchar(40)", "商品ID", "可空"),
            ("content", "text", "反馈内容", "非空"),
            ("status", "varchar(20)", "处理状态", "待处理/已处理"),
        ]),
        simple_schema_table("chat_sessions", [
            ("id", "varchar(40)", "会话ID", "非空"),
            ("user_id", "varchar(40)", "用户ID", "关联 users"),
            ("product_id", "varchar(40)", "商品ID", "可空"),
            ("title", "varchar(100)", "会话标题", "可空"),
            ("created_at", "datetime", "创建时间", "非空"),
        ]),
        simple_schema_table("chat_messages", [
            ("id", "varchar(40)", "消息ID", "非空"),
            ("session_id", "varchar(40)", "会话ID", "关联 chat_sessions"),
            ("role", "varchar(20)", "消息角色", "user/assistant"),
            ("content", "text", "消息内容", "非空"),
            ("sent_at", "datetime", "发送时间", "非空"),
        ]),
    ]

    for idx, rows in enumerate(schema_tables, start=2):
        if idx < len(doc.tables):
            fill_table(doc.tables[idx], rows)

    if len(doc.tables) > 17:
        fill_table(doc.tables[17], [
            ["类别", "项目", "配置/版本", "备注说明"],
            ["运行环境", "操作系统", "Windows", "本地开发与测试环境"],
            ["前端框架", "uni-app", "Vue 语法", "用于小程序页面开发"],
            ["后端框架", "Express", "Node.js", "提供 RESTful API"],
            ["数据库", "MySQL", "关系型数据库", "保存业务数据"],
            ["AI 服务", "大模型接口", "后端代理调用", "用于智能客服"],
        ])

    if len(doc.tables) > 18:
        fill_table(doc.tables[18], [
            ["用例编号", "测试项", "前置条件", "操作步骤", "预期结果", "实测结果"],
            ["TC-01", "用户登录与鉴权", "存在有效用户", "登录页提交身份信息", "登录成功并保存用户状态", "通过"],
            ["TC-02", "商品列表加载", "后端服务正常", "进入首页", "商品列表正确展示", "通过"],
            ["TC-03", "购物车管理", "用户已登录", "加入商品并修改数量", "购物车数量和金额同步变化", "通过"],
            ["TC-04", "订单创建", "购物车存在商品", "提交订单", "生成订单并清理购物车", "通过"],
            ["TC-05", "AI 智能客服", "用户已登录", "发送商品相关问题", "返回商品上下文回答", "通过"],
        ])


def main() -> None:
    doc = Document(str(TEMPLATE))

    update_tables(doc)

    # cover / English title / declaration dates
    fixed_paragraphs = {
        11: "二〇二六 年 五 月 六 日",
        18: TITLE_EN,
        33: "Date: 2026-05-06",
        43: "                              签字日期： 2026年 5月 6日",
        48: "                              签字日期： 2026年 5月 6日",
    }
    for idx, text in fixed_paragraphs.items():
        if idx < len(doc.paragraphs):
            doc.paragraphs[idx].text = text

    body = build_body()

    # Remove the old thesis content from "摘要" onward. This keeps the cover,
    # declaration pages, page setup and document styles from the template.
    start_element = doc.paragraphs[52]._element
    parent = start_element.getparent()
    children = list(parent)
    start_index = children.index(start_element)
    for child in children[start_index:]:
        parent.remove(child)

    for level, text in body:
        if text == TITLE_CN:
            continue
        if level == 0:
            doc.add_paragraph(text)
        elif level <= 2:
            doc.add_paragraph(text, style="Heading 1")
        elif level == 3:
            doc.add_paragraph(text, style="Heading 2")
        else:
            doc.add_paragraph(text, style="Heading 3")

    doc.save(str(OUTPUT))
    print(str(OUTPUT))


if __name__ == "__main__":
    main()
